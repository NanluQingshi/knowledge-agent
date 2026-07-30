"""DOCX 加载器 — 支持 .docx 文件."""

from __future__ import annotations

from pathlib import Path

from knowledge_agent.loaders.base import BaseLoader, Document


class DocxLoader(BaseLoader):
    """Word 文档加载器，支持 .docx 格式.

    依赖 python-docx 包，未安装时会给出清晰的安装提示。
    """

    def can_handle(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == ".docx"

    def load(self, file_path: Path) -> list[Document]:
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError(
                "python-docx is required to load .docx files. "
                "Install it with: pip install python-docx"
            )

        doc = DocxDocument(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = "\n".join(paragraphs)

        if not content.strip():
            return []

        stat = file_path.stat()
        return [
            Document(
                content=content,
                metadata={
                    "filename": file_path.name,
                    "size": stat.st_size,
                    "type": "docx",
                    "paragraphs": len(paragraphs),
                },
                source=str(file_path.resolve()),
            )
        ]
